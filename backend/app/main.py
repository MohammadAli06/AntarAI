"""
AntarAI Backend — Sovereign On-Premise Agentic AI Workbench
===========================================================

FastAPI application serving the core API for the AntarAI workbench.
Designed for MRPL (Mangalore Refinery) — Smart India Hackathon PS 26117.

Run with:
    cd backend
    uvicorn app.main:app --reload --port 8000
"""

from __future__ import annotations

import json
import hashlib
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Optional

import yaml
from fastapi import Depends, FastAPI, File, Form, HTTPException, UploadFile, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.agent.orchestrator import run_agent, run_agent_stream
from app.auth import (
    DEMO_MODE,
    Principal,
    create_access_token,
    create_demo_token,
    get_current_user,
    require_role,
    verify_password,
)
from app.database import (
    Attachment,
    Conversation,
    Document,
    Message,
    ModelAdmission,
    SessionLocal,
    Task,
    TaskAccess,
    User,
    create_tables,
    get_db,
)
from app.models import registry as model_registry
from app.models import admission as model_admission
from app.models.registry import get_call_count, list_models
from app.models.schemas import ModelEntry
from app.system_log import get_recent, log_event
from app.tools.ocr_extractor import extract_text
from app.tools import registry as tool_registry
from app.rag.ingestor import delete_document, ingest_document, search_document, search_all_documents
from app.rag.seed_knowledge import seed_knowledge_if_empty
from app.sovereignty.inspector import (
    count_external_calls,
    get_blocked_attempts,
    model_integrity,
    probe_local_services,
)
from seed import seed_users

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

_BACKEND_ROOT = Path(__file__).resolve().parents[1]
_DATA_DIR = _BACKEND_ROOT / "data"
_DOCUMENTS_DIR = _DATA_DIR / "documents"
_IMAGES_DIR = _DATA_DIR / "images"
_OUTPUTS_DIR = _BACKEND_ROOT / "outputs"

for d in (_DOCUMENTS_DIR, _IMAGES_DIR, _OUTPUTS_DIR):
    d.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# App
# ---------------------------------------------------------------------------

app = FastAPI(
    title="AntarAI — Sovereign Agentic AI Workbench",
    description=(
        "On-premise, air-gapped AI workbench for MRPL. "
        "Powered by Qwen3-8B-Q4_K_M (Qwen/Qwen3-8B-GGUF) via llama.cpp. "
        "Zero external network calls by design."
    ),
    version="1.0.0",
)

# CORS — allow frontend dev server (e.g. localhost:3000, localhost:5173)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------------------------------------------------------
# Startup: DB tables + seed demo users
# ---------------------------------------------------------------------------

@app.on_event("startup")
def on_startup() -> None:
    create_tables()
    _repair_legacy_document_metadata()
    seed_users()
    seed_knowledge_if_empty()
    import random, string
    worker_id = "wkr-" + "".join(random.choices(string.ascii_lowercase + string.digits, k=8))
    log_event("INFO", f"Task worker initialized: {worker_id}")



# ---------------------------------------------------------------------------
# Pydantic schemas
# ---------------------------------------------------------------------------

class LoginRequest(BaseModel):
    username: str
    password: str


class LoginResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    username: str
    role: str


class ModelInspectRequest(BaseModel):
    model_config = {"protected_namespaces": ()}

    model_path: str


class ChatResponse(BaseModel):
    model_config = {"protected_namespaces": ()}
    response: str
    model_used: str
    steps: list[str]
    generated_file: Optional[str] = None


class SovereigntyStatus(BaseModel):
    external_calls: int = 0
    local_model_calls: int
    local_files_accessed: int
    verdict: str = "✅ SOVEREIGN — All processing remained on-premise"


class SwitchRoleRequest(BaseModel):
    role: str


class CreateConversationRequest(BaseModel):
    title: Optional[str] = None


class UpdateConversationRequest(BaseModel):
    title: Optional[str] = None
    archived: Optional[bool] = None


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _save_attached_file(
    file: Optional[UploadFile],
    current_user: Principal,
    db: Session,
) -> tuple[Optional[str], Optional[str], bool]:
    """Persist an uploaded file + log a Document row. Returns (path, name, has_file)."""
    if file is None or not file.filename:
        return None, None, False

    filename = Path(file.filename).name
    ext = Path(filename).suffix.lower()
    save_dir = (
        _IMAGES_DIR
        if ext in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"}
        else _DOCUMENTS_DIR
    )
    dest = save_dir / filename
    content = file.file.read()
    with open(dest, "wb") as f:
        f.write(content)

    doc = Document(
        filename=filename,
        file_type=ext.lstrip("."),
        size_bytes=len(content),
        uploaded_by=current_user.id,
        indexed="not_indexed",
        failure_reason="Task attachment; not added to the shared Knowledge Base",
    )
    db.add(doc)
    db.commit()
    return str(dest), filename, True


def _knowledge_file_path(doc: Document) -> Optional[Path]:
    """Resolve a document's stored path while keeping it inside data folders."""
    name = doc.stored_filename or doc.filename
    if not name or Path(name).name != name:
        return None
    base = _IMAGES_DIR if (doc.file_type or "").lower() in {
        "png", "jpg", "jpeg", "gif", "bmp", "tiff", "webp"
    } else _DOCUMENTS_DIR
    candidate = (base / name).resolve()
    if base.resolve() not in candidate.parents:
        return None
    return candidate


def _index_knowledge_document(doc: Document, path: Path, db: Session) -> dict:
    """Move a Knowledge Base row to an explicit indexed/failed terminal state."""
    chunks = 0
    reason: Optional[str] = None
    try:
        if not tool_registry.is_tool_enabled("ocr"):
            raise RuntimeError("OCR Engine disabled by administrator")
        extracted = extract_text(str(path))
        if not extracted or extracted.lstrip().startswith("["):
            reason = extracted.strip("[] ") if extracted else "No readable text was extracted"
            raise RuntimeError(reason)
        result = ingest_document(text=extracted, filename=doc.filename, doc_id=doc.id)
        chunks = int(result.get("chunks", 0))
        if result.get("status") != "indexed" or chunks < 1:
            reason = result.get("reason") or "Vector index did not accept any chunks"
            raise RuntimeError(reason)
        doc.indexed = "indexed"
        doc.chunks_indexed = chunks
        doc.failure_reason = None
    except Exception as exc:
        doc.indexed = "failed"
        doc.chunks_indexed = 0
        doc.failure_reason = reason or str(exc) or "Indexing failed"
    db.commit()
    db.refresh(doc)
    return {
        "indexed": doc.indexed,
        "chunks_indexed": doc.chunks_indexed or 0,
        "failure_reason": doc.failure_reason,
    }


def _repair_legacy_document_metadata() -> None:
    """Backfill hashes and replace stale legacy pending/duplicate states."""
    session = SessionLocal()
    try:
        seen: dict[str, int] = {}
        for doc in session.query(Document).order_by(Document.id.asc()).all():
            path = _knowledge_file_path(doc)
            if path and path.is_file():
                digest = doc.content_hash or hashlib.sha256(path.read_bytes()).hexdigest()
                doc.content_hash = digest
                canonical_id = seen.get(digest)
                if canonical_id is not None and canonical_id != doc.id:
                    doc.indexed = "duplicate"
                    doc.chunks_indexed = 0
                    doc.failure_reason = f"Duplicate of document #{canonical_id}"
                    continue
                seen[digest] = doc.id
            if doc.indexed in {None, "pending", "processing", "unavailable"}:
                doc.indexed = "failed"
                doc.chunks_indexed = 0
                doc.failure_reason = (
                    doc.failure_reason
                    or "Legacy upload did not complete indexing; use Retry indexing"
                )
        session.commit()
    except Exception as exc:
        session.rollback()
        log_event("WARN", f"Knowledge metadata repair skipped: {exc}")
    finally:
        session.close()


def _persist_task_update(task_id: int, **fields) -> None:
    """Update a Task row from within the streaming generator (own session)."""
    s = SessionLocal()
    try:
        t = s.query(Task).filter(Task.id == task_id).first()
        if t is not None:
            for key, value in fields.items():
                if value is not None:
                    setattr(t, key, value)
            s.commit()
    finally:
        s.close()


def _conversation_or_404(db: Session, conv_id: int, current_user: Principal) -> Conversation:
    conv = db.query(Conversation).filter(Conversation.id == conv_id).first()
    if not conv:
        raise HTTPException(status_code=404, detail=f"Conversation #{conv_id} not found")
    if conv.user_id != current_user.id and current_user.role not in {"admin"}:
        raise HTTPException(status_code=403, detail="Not your conversation")
    if conv.deleted_at is not None:
        raise HTTPException(status_code=410, detail="Conversation deleted")
    return conv


def _conversation_to_dict(conv: Conversation, message_count: Optional[int] = None, last_message: Optional[str] = None) -> dict:
    preview = (conv.title or "New conversation").strip() or "New conversation"
    return {
        "id": conv.id,
        "user_id": conv.user_id,
        "title": preview,
        "created_at": conv.created_at.isoformat() if conv.created_at else None,
        "updated_at": conv.updated_at.isoformat() if conv.updated_at else None,
        "archived": bool(conv.archived),
        "message_count": message_count,
        "last_message_preview": (last_message or "")[:160] if last_message is not None else None,
    }


def _message_to_dict(m: Message, attachments: Optional[list[Attachment]] = None) -> dict:
    return {
        "id": m.id,
        "conversation_id": m.conversation_id,
        "role": m.role,
        "content": m.content,
        "task_id": m.task_id,
        "created_at": m.created_at.isoformat() if m.created_at else None,
        "attachments": [
            {"id": a.id, "filename": a.filename, "file_type": a.file_type, "file_path": a.file_path, "size_bytes": a.size_bytes}
            for a in (attachments or [])
        ],
    }


def _task_visible_to(task: Task, viewer: Principal, db: Session) -> bool:
    if task.user_id == viewer.id:
        return True
    if viewer.role in {"admin", "approver"}:
        if task.status in {"pending_approval", "approved", "rejected"}:
            return True
        has_grant = db.query(TaskAccess).filter(TaskAccess.task_id == task.id, TaskAccess.user_id == viewer.id).first() is not None
        if has_grant:
            return True
        if viewer.role == "admin":
            return True
    return False


def _build_conversation_history(db: Session, conversation_id: int, limit: int = 12) -> list[dict]:
    """Load prior messages (with file names) to thread into the model prompt."""
    rows = (
        db.query(Message)
        .filter(Message.conversation_id == conversation_id)
        .order_by(Message.created_at.asc())
        .all()
    )
    if not rows:
        return []
    rows = rows[-limit:]
    history: list[dict] = []
    for m in rows:
        atts = db.query(Attachment).filter(Attachment.message_id == m.id).all()
        history.append({
            "role": m.role,
            "content": m.content or "",
            "attachments": [a.filename for a in atts],
        })
    return history


def _auto_title(text: str) -> str:
    t = " ".join((text or "").strip().split())
    if not t:
        return "New conversation"
    return t[:72]


def _task_to_dict(t: Task, owner_name: Optional[str] = None) -> dict:
    """Serialise a Task row, including provenance + verification fields."""
    verification = None
    if t.verification_json:
        try:
            verification = json.loads(t.verification_json)
        except Exception:
            verification = None
    approval = None
    if t.status == "approved" and t.approved_by:
        approval = {
            "approvedBy": t.approved_by,
            "approvedAt": t.approved_at.isoformat() if t.approved_at else None,
            "taskId": f"TASK-{t.id}",
            "artifactHash": t.artifact_sha256 or "",
            "modelRunId": t.model_run_id or "",
            "evidenceSetId": f"EV-{t.id}-{t.evidence_count or 0}",
        }
    return {
        "id": t.id,
        "user_id": t.user_id,
        "task_type": t.task_type,
        "model_used": t.model_used,
        "prompt_preview": t.prompt_preview,
        "prompt_text": t.prompt_text,
        "input_filename": t.input_filename,
        "final_output": t.final_output,
        "generated_file": t.generated_file,
        "status": t.status,
        "timestamp": t.timestamp.isoformat() if t.timestamp else None,
        "risk": t.risk,
        "evidence_count": t.evidence_count,
        "artifact_sha256": t.artifact_sha256,
        "model_run_id": t.model_run_id,
        "verification": verification,
        "owner_name": owner_name,
        "approval": approval,
    }


# ---------------------------------------------------------------------------
# Public routes (no auth required)
# ---------------------------------------------------------------------------

@app.get("/", tags=["health"])
async def root():
    """Health check."""
    return {
        "service": "AntarAI Backend",
        "status": "running",
        "sovereignty": "enforced — zero external calls",
        "version": "0.2.0",
    }


@app.post("/auth/login", response_model=LoginResponse, tags=["auth"])
async def login(body: LoginRequest, db: Session = Depends(get_db)):
    """
    Authenticate with username + password.
    Returns a Bearer JWT valid for 8 hours.
    """
    user = db.query(User).filter(User.username == body.username).first()
    if not user or not verify_password(body.password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid username or password",
        )

    token = create_access_token(
        username=user.username,
        role=user.role,
        user_id=user.id,
    )
    return LoginResponse(
        access_token=token,
        username=user.username,
        role=user.role,
    )


# ---------------------------------------------------------------------------
# Protected routes — require valid JWT
# ---------------------------------------------------------------------------

@app.post("/chat", response_model=ChatResponse, tags=["chat"])
async def chat(
    message: str = Form(...),
    file: Optional[UploadFile] = File(None),
    conversation_id: Optional[int] = Form(None),
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Main chat endpoint (legacy JSON). If conversation_id is supplied it must be
    owned by the caller; otherwise a new conversation is created for continuity.
    """
    conv = None
    history: list[dict] = []
    if conversation_id is not None:
        conv = _conversation_or_404(db, conversation_id, current_user)
        history = _build_conversation_history(db, conv.id)

    file_path: Optional[str] = None
    filename: Optional[str] = None
    has_file = file is not None and file.filename

    if has_file:
        filename = Path(file.filename).name
        ext = Path(filename).suffix.lower()
        save_dir = (
            _IMAGES_DIR
            if ext in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"}
            else _DOCUMENTS_DIR
        )
        dest = save_dir / filename
        with open(dest, "wb") as f:
            f.write(await file.read())
        file_path = str(dest)
        doc = Document(
            filename=filename,
            file_type=ext.lstrip("."),
            size_bytes=dest.stat().st_size,
            uploaded_by=current_user.id,
            indexed="not_indexed",
            failure_reason="Task attachment; not added to the shared Knowledge Base",
        )
        db.add(doc)
        db.commit()

    if conv is None:
        conv = Conversation(user_id=current_user.id, title=_auto_title(message))
        db.add(conv)
        db.commit()
        db.refresh(conv)

    conv.updated_at = datetime.utcnow()
    db.add(Message(conversation_id=conv.id, role="user", content=message))
    db.commit()

    result = run_agent(
        message=message,
        has_file=bool(has_file),
        filename=filename,
        file_path=file_path,
        history=history,
    )

    if file_path and filename:
        last_user = db.query(Message).filter(Message.conversation_id == conv.id, Message.role == "user").order_by(Message.id.desc()).first()
        db.add(Attachment(conversation_id=conv.id, message_id=last_user.id if last_user else None, file_path=file_path, filename=filename, file_type=Path(filename).suffix.lstrip("."), size_bytes=Path(file_path).stat().st_size if Path(file_path).exists() else None))
        db.commit()

    initial_status = "pending_approval" if result.generated_file else "approved"
    task = Task(
        user_id=current_user.id,
        conversation_id=conv.id,
        task_type=result.model_used,
        model_used=result.model_used,
        prompt_preview=message[:120],
        prompt_text=message,
        input_filename=filename,
        final_output=result.response,
        generated_file=result.generated_file,
        status=initial_status,
    )
    db.add(task)
    db.commit()
    db.refresh(task)

    assistant_msg = Message(conversation_id=conv.id, role="assistant", content=result.response, task_id=task.id)
    db.add(assistant_msg)
    db.commit()
    if conv.title == "New conversation" and message.strip():
        conv.title = _auto_title(message)
        db.commit()

    return ChatResponse(
        response=result.response,
        model_used=result.model_used,
        steps=result.steps,
        generated_file=result.generated_file,
    )


# ---------------------------------------------------------------------------
# Streaming chat — real SSE pipeline (text/event-stream)
# ---------------------------------------------------------------------------

@app.post("/chat/stream", tags=["chat"])
async def chat_stream(
    message: str = Form(...),
    file: Optional[UploadFile] = File(None),
    conversation_id: Optional[int] = Form(None),
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Stream the agentic pipeline as Server-Sent Events.

    conversation_id — when supplied, must belong to the caller and the prior
    turns (including file references) are threaded into the model prompt so a
    follow-up like "what about the left side?" sees the same image. Auth +
    conversation-ownership is enforced BEFORE any model work.
    """
    conv = None
    history: list[dict] = []
    if conversation_id is not None:
        conv = _conversation_or_404(db, conversation_id, current_user)
        history = _build_conversation_history(db, conv.id)
    else:
        conv = Conversation(user_id=current_user.id, title=_auto_title(message))
        db.add(conv)
        db.commit()
        db.refresh(conv)

    file_path, filename, has_file = _save_attached_file(file, current_user, db)

    # Persist the user turn up-front so history is complete even if the stream is cancelled.
    user_msg = Message(conversation_id=conv.id, role="user", content=message)
    db.add(user_msg)
    db.commit()
    db.refresh(user_msg)
    if file_path and filename:
        db.add(Attachment(conversation_id=conv.id, message_id=user_msg.id, file_path=file_path, filename=filename, file_type=Path(filename).suffix.lstrip("."), size_bytes=Path(file_path).stat().st_size if Path(file_path).exists() else None))
        db.commit()
    if conv.title == "New conversation" and message.strip():
        conv.title = _auto_title(message)
        conv.updated_at = datetime.utcnow()
        db.commit()
    db.refresh(conv)
    conv_id_value = conv.id

    task = Task(
        user_id=current_user.id,
        conversation_id=conv_id_value,
        task_type="general",
        model_used="pending",
        prompt_preview=message[:120],
        prompt_text=message,
        input_filename=filename,
        status="planning",
    )
    db.add(task)
    db.commit()
    db.refresh(task)
    task_db_id = task.id
    task_id = f"TASK-{task.id}"

    def event_generator():
        final_response: Optional[str] = None
        final_status_holder: list[Optional[str]] = [None]
        for ev in run_agent_stream(message, has_file, filename, file_path, history=history):
            etype = ev["type"]
            data = ev.get("data", {})

            if etype == "router.completed":
                _persist_task_update(
                    task_db_id,
                    task_type=data.get("role", "general"),
                    model_used=data.get("model", ""),
                    status="running",
                    risk=data.get("risk"),
                )
            elif etype == "verification.started":
                _persist_task_update(task_db_id, status="verifying")
            elif etype == "task.completed":
                final_response = data.get("response")
                final_status_holder[0] = data.get("status", "completed")
                _persist_task_update(
                    task_db_id,
                    status=data.get("status", "completed"),
                    generated_file=data.get("generatedFile"),
                    risk=data.get("risk"),
                    evidence_count=data.get("evidenceCount"),
                    model_run_id=data.get("modelRunId"),
                    artifact_sha256=data.get("artifactSha256"),
                    verification_json=json.dumps(data.get("verification")),
                    final_output=data.get("response"),
                )
            elif etype == "task.failed":
                final_status_holder[0] = "failed"
                _persist_task_update(task_db_id, status="failed")

            payload = {
                "type": etype,
                "taskId": task_id,
                "conversationId": conv_id_value,
                "stepId": ev.get("stepId"),
                "timestamp": datetime.utcnow().isoformat() + "Z",
                "data": data,
            }
            yield f"data: {json.dumps(payload)}\n\n"

        if final_response is not None:
            s2 = SessionLocal()
            try:
                s2.add(Message(conversation_id=conv_id_value, role="assistant", content=final_response, task_id=task_db_id))
                c2 = s2.query(Conversation).filter(Conversation.id == conv_id_value).first()
                if c2:
                    c2.updated_at = datetime.utcnow()
                s2.commit()
            finally:
                s2.close()

        yield f"data: {json.dumps({'type': 'stream.end', 'taskId': task_id, 'conversationId': conv_id_value, 'timestamp': datetime.utcnow().isoformat() + 'Z'})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ---------------------------------------------------------------------------
# Conversations — user-owned history (auth + ownership enforced server-side)
# ---------------------------------------------------------------------------

@app.get("/conversations", tags=["conversations"])
async def list_conversations(
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
    q: Optional[str] = None,
    limit: int = 100,
    include_archived: bool = False,
):
    """List conversations belonging to the caller. No role-based bypass."""
    query = db.query(Conversation).filter(
        Conversation.user_id == current_user.id,
        Conversation.deleted_at.is_(None),
    )
    if not include_archived:
        query = query.filter(Conversation.archived.is_(False))
    if q and q.strip():
        like = f"%{q.strip()}%"
        query = query.filter(Conversation.title.ilike(like))
    convs = query.order_by(Conversation.updated_at.desc()).limit(limit).all()
    result = []
    for c in convs:
        last = db.query(Message).filter(Message.conversation_id == c.id).order_by(Message.id.desc()).first()
        count = db.query(Message).filter(Message.conversation_id == c.id).count()
        result.append(_conversation_to_dict(c, message_count=count, last_message=last.content if last else None))
    return {"conversations": result}


@app.post("/conversations", tags=["conversations"])
async def create_conversation(
    body: CreateConversationRequest,
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    title = (body.title or "New conversation").strip() or "New conversation"
    conv = Conversation(user_id=current_user.id, title=title[:120])
    db.add(conv)
    db.commit()
    db.refresh(conv)
    return {"conversation": _conversation_to_dict(conv, message_count=0, last_message=None)}


@app.get("/conversations/{conversation_id}", tags=["conversations"])
async def get_conversation(
    conversation_id: int,
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    conv = _conversation_or_404(db, conversation_id, current_user)
    messages = db.query(Message).filter(Message.conversation_id == conv.id).order_by(Message.created_at.asc()).all()
    tasks = {t.id: t for t in db.query(Task).filter(Task.conversation_id == conv.id).all()}
    enriched = []
    for m in messages:
        atts = db.query(Attachment).filter(Attachment.message_id == m.id).all()
        task_dict = _task_to_dict(tasks[m.task_id]) if m.task_id and m.task_id in tasks else None
        enriched.append({**_message_to_dict(m, atts), "task": task_dict})
    atts_all = db.query(Attachment).filter(Attachment.conversation_id == conv.id).all()
    return {
        "conversation": _conversation_to_dict(conv, message_count=len(messages), last_message=messages[-1].content if messages else None),
        "messages": enriched,
        "attachments": [
            {"id": a.id, "filename": a.filename, "file_type": a.file_type, "file_path": a.file_path, "size_bytes": a.size_bytes, "message_id": a.message_id}
            for a in atts_all
        ],
    }


@app.patch("/conversations/{conversation_id}", tags=["conversations"])
async def update_conversation(
    conversation_id: int,
    body: UpdateConversationRequest,
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    conv = _conversation_or_404(db, conversation_id, current_user)
    if body.title is not None:
        title = body.title.strip()
        if not title:
            raise HTTPException(status_code=400, detail="Title cannot be empty")
        conv.title = title[:120]
    if body.archived is not None:
        conv.archived = bool(body.archived)
    conv.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(conv)
    return {"conversation": _conversation_to_dict(conv)}


@app.delete("/conversations/{conversation_id}", tags=["conversations"])
async def delete_conversation(
    conversation_id: int,
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    conv = _conversation_or_404(db, conversation_id, current_user)
    conv.deleted_at = datetime.utcnow()
    conv.archived = True
    db.commit()
    return {"status": "deleted", "conversation_id": conv.id}


@app.get("/conversations/{conversation_id}/tasks", tags=["conversations"])
async def list_conversation_tasks(
    conversation_id: int,
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    conv = _conversation_or_404(db, conversation_id, current_user)
    tasks = db.query(Task).filter(Task.conversation_id == conv.id).order_by(Task.timestamp.asc()).all()
    return {"conversation_id": conv.id, "tasks": [_task_to_dict(t) for t in tasks]}


# ---------------------------------------------------------------------------
# Task detail + governed visibility: owner always; approver/admin via status/access
# ---------------------------------------------------------------------------

@app.get("/tasks/{task_id}", tags=["history"])
async def get_task_detail(
    task_id: int,
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    task = db.query(Task).filter(Task.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail=f"Task #{task_id} not found")
    if not _task_visible_to(task, current_user, db):
        raise HTTPException(status_code=403, detail="Not authorized to view this task")
    owner = db.query(User).filter(User.id == task.user_id).first()
    detail = _task_to_dict(task, owner_name=owner.username if owner else None)
    # Attach conversation slice (still governed by _task_visible_to above)
    conv_msgs: list[dict] = []
    if task.conversation_id:
        for m in db.query(Message).filter(Message.conversation_id == task.conversation_id).order_by(Message.created_at.asc()).all():
            conv_msgs.append(_message_to_dict(m, db.query(Attachment).filter(Attachment.message_id == m.id).all()))
    return {"task": detail, "messages": conv_msgs if task.conversation_id else []}


@app.post("/tasks/{task_id}/share", tags=["approvals"])
async def share_task_for_review(
    task_id: int,
    body: dict,
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Owner shares a governed task with an approver (task-level controlled sharing)."""
    task = db.query(Task).filter(Task.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail=f"Task #{task_id} not found")
    if task.user_id != current_user.id and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only the task owner (or admin) can share it")
    target_username = str(body.get("username") or body.get("user") or "").strip()
    if not target_username:
        raise HTTPException(status_code=400, detail="username is required")
    target = db.query(User).filter(User.username == target_username).first()
    if not target:
        raise HTTPException(status_code=404, detail=f"User '{target_username}' not found")
    if target.role not in {"approver", "admin"}:
        raise HTTPException(status_code=400, detail="Can only share with approver/admin roles")
    existing = db.query(TaskAccess).filter(TaskAccess.task_id == task.id, TaskAccess.user_id == target.id).first()
    if not existing:
        db.add(TaskAccess(task_id=task.id, user_id=target.id, access_type=str(body.get("access_type") or "review")))
        db.commit()
    return {"status": "shared", "task_id": task.id, "shared_with": target.username}


@app.get("/models", tags=["models"])
async def get_models(current_user: Principal = Depends(get_current_user)):
    """Return list of configured models and their current status."""
    return {"models": list_models()}


@app.post("/admin/models", tags=["admin"])
async def add_registered_model(
    entry: ModelEntry,
    current_user: Principal = Depends(require_role(["admin"])),
):
    try:
        model_registry.add_model(entry)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return {"status": "added", "models": model_registry.list_models()}


@app.get("/admin/models/files", tags=["admin"])
async def list_local_model_files(
    current_user: Principal = Depends(require_role(["admin"])),
):
    """List GGUF files placed on this node for metadata-driven registration."""
    return {"files": model_registry.list_model_files()}


@app.post("/admin/models/inspect", tags=["admin"])
async def inspect_local_model_file(
    body: ModelInspectRequest,
    current_user: Principal = Depends(require_role(["admin"])),
):
    try:
        return {"metadata": model_registry.inspect_model_file(body.model_path)}
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.delete("/admin/models/{role}", tags=["admin"])
async def remove_registered_model(
    role: str,
    current_user: Principal = Depends(require_role(["admin"])),
):
    try:
        model_registry.remove_model(role)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=f"No model registered for role: {role}") from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return {"status": "removed", "models": model_registry.list_models()}


@app.patch("/admin/models/{role}/endpoint", tags=["admin"])
async def update_registered_model_endpoint(
    role: str,
    body: dict,
    current_user: Principal = Depends(require_role(["admin"])),
):
    endpoint = str(body.get("endpoint", "")).strip()
    try:
        model_registry.update_model_endpoint(role, endpoint)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=f"No model registered for role: {role}") from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return {"status": "updated", "models": model_registry.list_models()}


@app.post("/admin/models/reload", tags=["admin"])
async def reload_registered_models(
    current_user: Principal = Depends(require_role(["admin"])),
):
    model_registry.reload_models()
    return {"status": "reloaded", "models": model_registry.list_models()}


# ---------------------------------------------------------------------------
# Model Center — sovereign model admission pipeline
# ---------------------------------------------------------------------------

class AdmissionPrecheckRequest(BaseModel):
    model_config = {"protected_namespaces": ()}

    source: str = "catalog"           # catalog | local | offline-package
    role: str = "general"
    catalog_key: Optional[str] = None
    model_path: str = ""


class AdmissionStartRequest(BaseModel):
    model_config = {"protected_namespaces": ()}

    source: str = "catalog"
    role: str = "general"
    catalog_key: Optional[str] = None
    model_path: str = ""
    description: str = ""
    capabilities: Optional[list] = None
    runtime_context_tokens: int = 4096


@app.get("/admin/models/catalog", tags=["admin"])
async def model_catalog(current_user: Principal = Depends(require_role(["admin"]))):
    """Curated sovereign catalog + locally available GGUF packages."""
    return {
        "catalog": model_admission.catalog_entries(),
        "local_files": model_registry.list_model_files(),
        "policy": (yaml.safe_load((_BACKEND_ROOT / "policies.yaml").read_text(encoding="utf-8")) or {}).get("model_admission", {}),
    }


@app.post("/admin/models/admission/precheck", tags=["admin"])
async def admission_precheck(
    body: AdmissionPrecheckRequest,
    current_user: Principal = Depends(require_role(["admin"])),
):
    """Run the admission gate synchronously — the wizard's fit-check screen."""
    return model_admission.precheck(
        source=body.source,
        role=body.role,
        catalog_key=body.catalog_key,
        model_path=body.model_path,
    )


@app.post("/admin/models/admission/stream", tags=["admin"])
async def admission_stream(
    body: AdmissionStartRequest,
    current_user: Principal = Depends(require_role(["admin"])),
):
    """Run the full Model Admission Pipeline as Server-Sent Events."""

    def persist_audit(summary: dict) -> Optional[int]:
        session = SessionLocal()
        try:
            record = ModelAdmission(
                model_name=summary["model_name"],
                catalog_key=summary.get("catalog_key") or None,
                source=summary.get("source", "local"),
                role=summary.get("role", "general"),
                sha256=summary.get("sha256", ""),
                node=summary.get("node"),
                port=summary.get("port"),
                checks_json=json.dumps(summary.get("checks", [])),
                metadata_json=json.dumps(summary.get("metadata", {})),
                admitted_by=summary.get("admitted_by", current_user.username),
            )
            session.add(record)
            session.commit()
            session.refresh(record)
            return record.id
        finally:
            session.close()

    def event_generator():
        for ev in model_admission.admit_model_stream(
            source=body.source,
            role=body.role,
            catalog_key=body.catalog_key,
            model_path=body.model_path,
            description=body.description,
            capabilities=body.capabilities or [],
            runtime_context_tokens=body.runtime_context_tokens,
            admitted_by=current_user.username,
            persist_audit=persist_audit,
        ):
            yield f"data: {json.dumps(ev)}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.get("/admin/models/admissions", tags=["admin"])
async def admission_history(
    current_user: Principal = Depends(require_role(["admin"])),
    db: Session = Depends(get_db),
):
    """Audit trail of every model admitted through the Model Center."""
    rows = db.query(ModelAdmission).order_by(ModelAdmission.admitted_at.desc()).limit(50).all()
    result = []
    for row in rows:
        checks = None
        metadata = None
        if row.checks_json:
            try:
                checks = json.loads(row.checks_json)
            except ValueError:
                checks = None
        if row.metadata_json:
            try:
                metadata = json.loads(row.metadata_json)
            except ValueError:
                metadata = None
        result.append({
            "id": row.id,
            "model_name": row.model_name,
            "catalog_key": row.catalog_key,
            "source": row.source,
            "role": row.role,
            "sha256": row.sha256,
            "node": row.node,
            "port": row.port,
            "checks": checks,
            "metadata": metadata,
            "admitted_by": row.admitted_by,
            "admitted_at": row.admitted_at.isoformat() if row.admitted_at else None,
        })
    return {"admissions": result}


@app.post("/upload", tags=["files"])
async def upload_file(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Upload a file (PDF, image, etc.). Triggers local OCR + ChromaDB ingestion."""
    filename = Path(file.filename or "").name
    if not filename:
        raise HTTPException(status_code=400, detail="A filename is required")
    ext = Path(filename).suffix.lower()
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="The uploaded file is empty")
    content_hash = hashlib.sha256(content).hexdigest()
    existing = db.query(Document).filter(Document.content_hash == content_hash).first()
    if existing:
        return {
            "status": "duplicate",
            "existing_doc_id": existing.id,
            "filename": existing.filename,
            "indexed": existing.indexed,
            "chunks_indexed": existing.chunks_indexed or 0,
            "failure_reason": existing.failure_reason,
        }
    save_dir = (
        _IMAGES_DIR
        if ext in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"}
        else _DOCUMENTS_DIR
    )
    stored_filename = f"{content_hash[:16]}_{filename}"
    dest = save_dir / stored_filename
    with open(dest, "wb") as f:
        f.write(content)

    doc = Document(
        filename=filename,
        file_type=ext.lstrip("."),
        size_bytes=len(content),
        uploaded_by=current_user.id,
        indexed="processing",
        content_hash=content_hash,
        stored_filename=stored_filename,
        chunks_indexed=0,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    result = _index_knowledge_document(doc, dest, db)

    return {
        "status": result["indexed"],
        "doc_id": doc.id,
        "filename": filename,
        "size_bytes": len(content),
        **result,
    }



@app.get("/outputs", tags=["files"])
async def list_outputs(current_user: User = Depends(get_current_user)):
    """List all generated output files available for download."""
    files = []
    for f in _OUTPUTS_DIR.iterdir():
        if f.is_file():
            files.append(
                {
                    "filename": f.name,
                    "size_bytes": f.stat().st_size,
                    "download_url": f"/outputs/{f.name}",
                }
            )
    return {"outputs": files}


def _resolve_output_path(filename: str) -> Path:
    """Resolve an output filename without permitting directory traversal."""
    path = (_OUTPUTS_DIR / filename).resolve()
    if _OUTPUTS_DIR.resolve() not in path.parents or not path.is_file():
        raise HTTPException(status_code=404, detail=f"File '{filename}' not found")
    return path


@app.get("/outputs/{filename}/preview", tags=["files"])
async def preview_output(
    filename: str,
    current_user: User = Depends(get_current_user),
):
    """Return a safe text preview of the exact generated artifact."""
    path = _resolve_output_path(filename)
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            from docx import Document  # type: ignore
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif suffix in {".py", ".txt", ".md", ".json", ".csv", ".yaml", ".yml"}:
            text = path.read_text(encoding="utf-8", errors="replace")
        else:
            text = "Preview is unavailable for this file type. Download the artifact to inspect it."
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Could not preview '{filename}': {exc}") from exc
    return {"filename": path.name, "content": text[:50000], "truncated": len(text) > 50000}


@app.get("/outputs/{filename}", tags=["files"])
async def download_output(
    filename: str,
    current_user: User = Depends(get_current_user),
):
    """Download a generated output file."""
    file_path = _resolve_output_path(filename)
    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type="application/octet-stream",
    )


@app.get("/sovereignty-status", tags=["sovereignty"])
async def sovereignty_status(
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Real, measured sovereignty proof.

    Counts (all measured, none asserted):
      - external_calls   : live outbound ESTABLISHED sockets (non-loopback) via psutil
      - local_model_calls: real in-process counter incremented by call_model()
      - local_files_accessed: documents in SQLite
      - blocked_attempts : sandbox network interceptions
      - local_services   : real TCP port probes (backend + llama.cpp)
      - model_integrity  : SHA-256 of the on-prem weights (when configured)
    """
    model_calls = get_call_count()
    file_count = db.query(Document).count()
    outputs_count = sum(1 for f in _OUTPUTS_DIR.iterdir() if f.is_file())
    external = count_external_calls()
    services = probe_local_services()
    integrity = model_integrity()
    model_online = any(
        service.get("online") and "llama.cpp" in service.get("name", "")
        for service in services
    )

    # Policy compliance and runtime availability are separate conditions.
    if external > 0:
        verdict = "air_gap_violated"
    elif not model_online:
        verdict = "model_offline"
    else:
        verdict = "verified"

    return {
        "external_calls": external,
        "local_model_calls": model_calls,
        "local_files_accessed": file_count,
        "blocked_attempts": get_blocked_attempts(),
        "local_services": services,
        "model_integrity": integrity,
        "online": model_online,
        "verdict": verdict,
    }



# ---------------------------------------------------------------------------
# System log — live ring buffer feed for AdminHome terminal
# ---------------------------------------------------------------------------

@app.get("/system-log", tags=["admin"])
async def system_log(
    current_user: Principal = Depends(get_current_user),
    limit: int = 50,
):
    """Return the most recent system log entries from the in-process ring buffer.

    No elevated role required — same visibility scope as /sovereignty-status.
    """
    return {"entries": get_recent(limit)}


# ---------------------------------------------------------------------------
# System metrics — CPU, RAM, GPU/VRAM (orchestrator node only, Option A)
# ---------------------------------------------------------------------------

def _get_system_metrics() -> dict:
    """Read live CPU%, RAM%, and GPU/VRAM from the orchestrator node.

    psutil covers CPU and RAM (already a dependency).
    pynvml covers GPU/VRAM. On Windows NVIDIA systems without that Python
    package, use the locally installed ``nvidia-smi`` driver utility instead.
    Labelled as 'Orchestrator node' — honest about single-node scope.
    """
    cpu_percent: float = 0.0
    ram_used: int = 0
    ram_total: int = 1
    gpu_percent: float = 0.0
    vram_used: int = 0
    vram_total: int = 1
    gpu_available = False
    gpu_name = "N/A"

    try:
        import psutil  # type: ignore
        cpu_percent = psutil.cpu_percent(interval=0.1)
        mem = psutil.virtual_memory()
        ram_used = mem.used
        ram_total = mem.total
    except Exception:
        pass

    try:
        import pynvml  # type: ignore
        pynvml.nvmlInit()
        handle = pynvml.nvmlDeviceGetHandleByIndex(0)
        info = pynvml.nvmlDeviceGetMemoryInfo(handle)
        util = pynvml.nvmlDeviceGetUtilizationRates(handle)
        gpu_name = pynvml.nvmlDeviceGetName(handle)
        if isinstance(gpu_name, bytes):
            gpu_name = gpu_name.decode("utf-8", errors="replace")
        gpu_percent = float(util.gpu)
        vram_used = info.used
        vram_total = info.total
        gpu_available = True
        pynvml.nvmlShutdown()
    except Exception:
        # ``nvidia-smi`` ships with the NVIDIA display driver, so this keeps
        # the overview live without requiring a separate pynvml installation.
        # It is a local process query and makes no network call.
        try:
            result = subprocess.run(
                [
                    "nvidia-smi",
                    "--query-gpu=name,utilization.gpu,memory.used,memory.total",
                    "--format=csv,noheader,nounits",
                ],
                capture_output=True,
                text=True,
                check=True,
                timeout=3,
            )
            values = [value.strip() for value in result.stdout.splitlines()[0].split(",")]
            gpu_name, utilization, memory_used, memory_total = values
            gpu_percent = float(utilization)
            vram_used = int(float(memory_used) * 1024 * 1024)
            vram_total = int(float(memory_total) * 1024 * 1024)
            gpu_available = vram_total > 0
        except (FileNotFoundError, IndexError, ValueError, subprocess.SubprocessError):
            gpu_available = False

    return {
        "node": "Orchestrator",
        "cpu_percent": round(cpu_percent, 1),
        "ram_used_bytes": ram_used,
        "ram_total_bytes": ram_total,
        "ram_percent": round(ram_used / ram_total * 100, 1) if ram_total else 0,
        "gpu_available": gpu_available,
        "gpu_name": gpu_name,
        "gpu_percent": round(gpu_percent, 1),
        "vram_used_bytes": vram_used,
        "vram_total_bytes": vram_total,
        "vram_percent": round(vram_used / vram_total * 100, 1) if vram_total else 0,
    }


@app.get("/system-metrics", tags=["admin"])
async def system_metrics(
    current_user: Principal = Depends(require_role(["admin"])),
):
    """Live CPU / RAM / GPU metrics from the orchestrator node."""
    return _get_system_metrics()




@app.get("/tasks/mine", tags=["history"])
async def get_my_tasks(
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
    limit: int = 20,
):
    """Return task history for the currently logged-in user."""
    tasks = (
        db.query(Task)
        .filter(Task.user_id == current_user.id)
        .order_by(Task.timestamp.desc())
        .limit(limit)
        .all()
    )
    return {"tasks": [_task_to_dict(t, owner_name=current_user.username) for t in tasks]}


@app.get("/tasks", tags=["history"])
async def list_all_tasks(
    current_user: Principal = Depends(require_role(["approver", "admin"])),
    db: Session = Depends(get_db),
    limit: int = 50,
):
    """List ALL users' tasks — accessible by Approvers and Admins only."""
    tasks = (
        db.query(Task)
        .order_by(Task.timestamp.desc())
        .limit(limit)
        .all()
    )
    user_ids = {t.user_id for t in tasks}
    users = {u.id: u.username for u in db.query(User).filter(User.id.in_(user_ids)).all()} if user_ids else {}
    return {
        "tasks": [
            _task_to_dict(t, owner_name=users.get(t.user_id)) for t in tasks
        ]
    }


@app.get("/audit", tags=["history"])
async def audit_trail(
    current_user: Principal = Depends(require_role(["approver", "admin"])),
    db: Session = Depends(get_db),
    limit: int = 50,
):
    """Audit trail — recent task transitions (status, model, approver)."""
    tasks = db.query(Task).order_by(Task.timestamp.desc()).limit(limit).all()
    user_ids = {t.user_id for t in tasks}
    users = {u.id: u.username for u in db.query(User).filter(User.id.in_(user_ids)).all()} if user_ids else {}
    return {
        "events": [
            {
                "id": t.id,
                "taskId": f"TASK-{t.id}",
                "owner": users.get(t.user_id, "unknown"),
                "modelUsed": t.model_used,
                "status": t.status,
                "risk": t.risk,
                "evidenceCount": t.evidence_count,
                "approvedBy": t.approved_by,
                "approvedAt": t.approved_at.isoformat() if t.approved_at else None,
                "timestamp": t.timestamp.isoformat() if t.timestamp else None,
                "generatedFile": t.generated_file,
            }
            for t in tasks
        ]
    }


@app.post("/tasks/{task_id}/approve", tags=["approvals"])
async def approve_task(
    task_id: int,
    current_user: Principal = Depends(require_role(["approver", "admin"])),
    db: Session = Depends(get_db),
):
    """Approve an AI-generated draft/note — Approver and Admin only.

    Writes an immutable approval record (approver, timestamp, artifact hash,
    model run id, evidence set id) for provenance.
    """
    task = db.query(Task).filter(Task.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail=f"Task #{task_id} not found")

    task.status = "approved"
    task.approved_by = current_user.username
    task.approved_at = datetime.utcnow()
    db.commit()

    approval = {
        "approvedBy": current_user.username,
        "approvedAt": task.approved_at.isoformat(),
        "taskId": f"TASK-{task.id}",
        "artifactHash": task.artifact_sha256 or "",
        "modelRunId": task.model_run_id or "",
        "evidenceSetId": f"EV-{task.id}-{task.evidence_count or 0}",
    }
    return {"status": "approved", "task_id": task_id, "approval": approval}


@app.post("/tasks/{task_id}/reject", tags=["approvals"])
async def reject_task(
    task_id: int,
    current_user: Principal = Depends(require_role(["approver", "admin"])),
    db: Session = Depends(get_db),
):
    """Reject an AI-generated draft/note — Approver and Admin only."""
    task = db.query(Task).filter(Task.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail=f"Task #{task_id} not found")

    task.status = "rejected"
    task.approved_by = current_user.username
    task.approved_at = datetime.utcnow()
    db.commit()
    return {"status": "rejected", "task_id": task_id, "rejected_by": current_user.username}


# ---------------------------------------------------------------------------
# Demo role switching (DEMO_MODE-gated, server-verified)
# ---------------------------------------------------------------------------

@app.get("/me", tags=["auth"])
async def me(current_user: Principal = Depends(get_current_user)):
    """Return the authoritative (signed-token) role + demo flag."""
    return {
        "username": current_user.username,
        "role": current_user.role,
        "demo": current_user.demo,
        "demoMode": DEMO_MODE,
    }


@app.post("/demo/switch-role", tags=["auth"])
async def switch_demo_role(
    body: SwitchRoleRequest,
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Switch the demo role by re-issuing a short-lived, demo-scoped JWT.

    Gated behind DEMO_MODE (default ON for the final round). The new role
    lives only in the signed token — it is NOT persisted to the User row.
    """
    if not DEMO_MODE:
        raise HTTPException(status_code=404, detail="Demo role switching is disabled")

    user = db.query(User).filter(User.username == current_user.username).first()
    if user is None:
        raise HTTPException(status_code=404, detail="User not found")

    token = create_demo_token(user, body.role)
    return {
        "access_token": token,
        "token_type": "bearer",
        "username": user.username,
        "role": body.role,
        "demo": True,
        "demoMode": True,
    }


# ---------------------------------------------------------------------------
# Admin control plane — Tools / Users / Policies
# ---------------------------------------------------------------------------

@app.get("/tools", tags=["admin"])
async def list_tools(current_user: Principal = Depends(get_current_user)):
    """Tool registry with real availability status."""
    return {"tools": tool_registry.list_tools()}


@app.post("/admin/tools/{name}/toggle", tags=["admin"])
async def toggle_registered_tool(
    name: str,
    enabled: bool,
    current_user: Principal = Depends(require_role(["admin"])),
):
    try:
        tool_registry.set_enabled(name, enabled, toggled_by=current_user.username)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=f"Tool not found: {name}") from exc
    log_event("GOV", f"Tool '{name}' {'enabled' if enabled else 'disabled'} by {current_user.username}")
    return {"status": "updated", "tools": tool_registry.list_tools()}


@app.post("/admin/tools/reload", tags=["admin"])
async def reload_registered_tools(
    current_user: Principal = Depends(require_role(["admin"])),
):
    tool_registry.reload_tools()
    return {"status": "reloaded", "tools": tool_registry.list_tools()}


@app.get("/users", tags=["admin"])
async def list_users(
    current_user: Principal = Depends(require_role(["admin"])),
    db: Session = Depends(get_db),
):
    """List all users — Admin only."""
    users = db.query(User).order_by(User.created_at.desc()).all()
    return {
        "users": [
            {
                "id": u.id,
                "username": u.username,
                "role": u.role,
                "createdAt": u.created_at.isoformat() if u.created_at else None,
            }
            for u in users
        ]
    }


@app.get("/policies", tags=["admin"])
async def list_policies(
    current_user: Principal = Depends(require_role(["admin"])),
):
    """Return governance policies from policies.yaml — Admin only."""
    policies_path = _BACKEND_ROOT / "policies.yaml"
    if not policies_path.exists():
        raise HTTPException(status_code=404, detail="policies.yaml not found")
    with open(policies_path, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f)
    return {"policies": data}


@app.get("/documents", tags=["knowledge"])
async def list_documents(
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List uploaded documents tracked in the SQLite documents table."""
    docs = db.query(Document).order_by(Document.upload_date.desc()).all()
    return {
        "documents": [
            {
                "id": d.id,
                "filename": d.filename,
                "file_type": d.file_type,
                "size_bytes": d.size_bytes,
                "indexed": d.indexed,
                "chunks_indexed": d.chunks_indexed or 0,
                "failure_reason": d.failure_reason,
                "upload_date": d.upload_date.isoformat() if d.upload_date else None,
            }
            for d in docs
        ]
    }


@app.get("/documents/search", tags=["knowledge"])
async def global_documents_search(
    query: str,
    current_user: Principal = Depends(get_current_user),
):
    """Global RAG vector search across all indexed Knowledge Base documents."""
    q = query.strip()
    if len(q) < 2:
        return {"query": q, "matches": []}
    return {"query": q, "matches": search_all_documents(q, n_results=20)}


@app.get("/documents/{doc_id}/search", tags=["knowledge"])
async def search_within_document(
    doc_id: int,
    query: str,
    current_user: Principal = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Search only the vector chunks belonging to one indexed document."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.indexed == "duplicate":
        raise HTTPException(status_code=409, detail=doc.failure_reason or "Duplicate document")
    if doc.indexed != "indexed":
        raise HTTPException(status_code=409, detail=f"Document is {doc.indexed}, not indexed")
    if len(query.strip()) < 2:
        raise HTTPException(status_code=400, detail="Search query must contain at least 2 characters")
    return {"doc_id": doc_id, "query": query.strip(), "matches": search_document(doc_id, query)}


@app.post("/documents/{doc_id}/reindex", tags=["knowledge"])
async def reindex_knowledge_document(
    doc_id: int,
    current_user: Principal = Depends(require_role(["admin"])),
    db: Session = Depends(get_db),
):
    """Retry extraction and indexing for a failed or legacy pending document."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.indexed == "duplicate":
        raise HTTPException(status_code=409, detail=doc.failure_reason or "Duplicate document")
    path = _knowledge_file_path(doc)
    if path is None or not path.is_file():
        doc.indexed = "failed"
        doc.failure_reason = "The original uploaded file is no longer present on disk"
        doc.chunks_indexed = 0
        db.commit()
        raise HTTPException(status_code=409, detail=doc.failure_reason)
    doc.indexed = "processing"
    doc.failure_reason = None
    db.commit()
    result = _index_knowledge_document(doc, path, db)
    return {"status": result["indexed"], "doc_id": doc.id, **result}


@app.delete("/knowledge-base/{doc_id}", tags=["knowledge"])
async def delete_knowledge_document(
    doc_id: int,
    current_user: Principal = Depends(require_role(["admin"])),
    db: Session = Depends(get_db),
):
    """Delete a document from Knowledge Base — Admin only."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    delete_document(doc_id)
    path = _knowledge_file_path(doc)
    if path and path.is_file():
        # Old rows may share an unhashed filename; retain the file while another
        # row still references it.
        shared = False
        if not doc.stored_filename:
            shared = db.query(Document).filter(
                Document.id != doc.id,
                Document.filename == doc.filename,
            ).first() is not None
        if not shared:
            path.unlink()
    db.delete(doc)
    db.commit()
    return {"status": "deleted", "doc_id": doc_id}
