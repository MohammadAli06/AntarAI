"""
OCR Extractor — extracts text from PDFs and images for Qwen analysis.

Dependencies:
  pip install pytesseract pillow pypdf2

Tesseract binary must be installed separately:
  Windows: https://github.com/UB-Mannheim/tesseract/wiki
  Set TESSERACT_CMD env var if not in PATH.

Graceful fallback: if Tesseract is not available, returns a warning string
so the rest of the pipeline still functions (Qwen sees the warning instead
of crashing).
"""

from __future__ import annotations

import os
import logging
import io
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Optional imports with graceful degradation
# ---------------------------------------------------------------------------

try:
    import pytesseract
    from PIL import Image

    _tesseract_cmd = os.getenv("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
    if os.path.exists(_tesseract_cmd):
        pytesseract.pytesseract.tesseract_cmd = _tesseract_cmd
    # Import success alone does not prove the local binary can execute.
    pytesseract.get_tesseract_version()
    _TESSERACT_AVAILABLE = True
except Exception:
    _TESSERACT_AVAILABLE = False
    logger.warning("pytesseract, Pillow, or the Tesseract binary is unavailable")

try:
    import PyPDF2
    _PYPDF_AVAILABLE = True
except ImportError:
    _PYPDF_AVAILABLE = False
    logger.warning("PyPDF2 not installed — PDF text extraction unavailable")

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — DOCX extraction unavailable")

try:
    from openpyxl import load_workbook
    _XLSX_AVAILABLE = True
except ImportError:
    _XLSX_AVAILABLE = False
    logger.warning("openpyxl not installed — XLSX extraction unavailable")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text(file_path: str) -> str:
    """
    Extract readable text from a PDF or image file.

    Strategy:
      1. PDF → try digital text extraction (PyPDF2) first
      2. PDF → if no text found, render pages as images and OCR each
      3. Image → OCR directly via Tesseract
      4. Fallback → return descriptive message so Qwen still receives context

    Parameters
    ----------
    file_path : str
        Absolute path to the uploaded file.

    Returns
    -------
    str
        Extracted text content, or a fallback message.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_from_pdf(path)
    elif suffix in {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp", ".gif"}:
        return _extract_from_image(path)
    elif suffix in {".txt", ".md", ".csv", ".json", ".yaml", ".yml"}:
        try:
            text = path.read_text(encoding="utf-8", errors="replace").strip()
            return text or f"[No readable text found in {path.name}]"
        except Exception as exc:
            return f"[Text extraction failed for {path.name}: {exc}]"
    elif suffix == ".docx":
        return _extract_from_docx(path)
    elif suffix == ".xlsx":
        return _extract_from_xlsx(path)
    else:
        return f"[File type {suffix} not supported for text extraction — filename: {path.name}]"


def _extract_from_docx(path: Path) -> str:
    """Extract paragraphs and table cells from a Word document."""
    if not _DOCX_AVAILABLE:
        return f"[DOCX extraction unavailable for {path.name} — python-docx is not installed]"
    try:
        document = DocxDocument(str(path))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    blocks.append(" | ".join(values))
        text = "\n\n".join(blocks).strip()
        return text or f"[No readable text found in {path.name}]"
    except Exception as exc:
        logger.warning("DOCX extraction failed for %s: %s", path.name, exc)
        return f"[DOCX extraction failed for {path.name}: {exc}]"


def _extract_from_xlsx(path: Path) -> str:
    """Extract non-empty spreadsheet cells with real sheet boundaries."""
    if not _XLSX_AVAILABLE:
        return f"[XLSX extraction unavailable for {path.name} — openpyxl is not installed]"
    try:
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        blocks: list[str] = []
        for sheet in workbook.worksheets:
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(value).strip() if value is not None else "" for value in row]
                if any(values):
                    rows.append(" | ".join(values))
            blocks.append(f"--- Sheet: {sheet.title} ---\n" + "\n".join(rows))
        workbook.close()
        text = "\n\n".join(blocks).strip()
        return text or f"[No readable cells found in {path.name}]"
    except Exception as exc:
        logger.warning("XLSX extraction failed for %s: %s", path.name, exc)
        return f"[XLSX extraction failed for {path.name}: {exc}]"


def _extract_from_pdf(path: Path) -> str:
    """Try digital text first, then OCR fallback."""
    text = ""

    # 1. Digital text extraction (fast — works on non-scanned PDFs)
    if _PYPDF_AVAILABLE:
        try:
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages_text = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages_text.append(f"--- Page {i+1} ---\n{page_text.strip()}")
                text = "\n\n".join(pages_text)
        except Exception as exc:
            logger.warning("PyPDF2 extraction failed: %s", exc)

    if text.strip():
        logger.info("PDF digital text extracted: %d chars from %s", len(text), path.name)
        return text

    # 2. OCR fallback for scanned PDFs
    if _TESSERACT_AVAILABLE:
        try:
            try:
                import fitz  # type: ignore
                pdf = fitz.open(str(path))
                images = []
                for page in pdf:
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    images.append(Image.open(io.BytesIO(pixmap.tobytes("png"))))
                pdf.close()
            except ImportError:
                # Retained for installations that already provide Poppler.
                from pdf2image import convert_from_path  # type: ignore
                images = convert_from_path(str(path), dpi=200)
            ocr_pages = []
            for i, img in enumerate(images):
                page_text = pytesseract.image_to_string(img, lang="eng").strip()
                if page_text:
                    ocr_pages.append(f"--- Page {i+1} (OCR) ---\n{page_text}")
            text = "\n\n".join(ocr_pages)
            if text.strip():
                logger.info("PDF OCR text extracted: %d chars", len(text))
                return text
        except ImportError:
            logger.warning("PyMuPDF/pdf2image not installed — cannot OCR scanned PDF pages")
        except Exception as exc:
            logger.warning("PDF OCR failed: %s", exc)

    # 3. Nothing worked
    return (
        f"[Could not extract text from {path.name}. "
        "It may be a scanned PDF requiring Tesseract + pdf2image. "
        "Please describe the document contents in your message.]"
    )


def _extract_from_image(path: Path) -> str:
    """OCR an image file via Tesseract."""
    if not _TESSERACT_AVAILABLE:
        return (
            f"[OCR unavailable for {path.name} — Tesseract not installed. "
            "Install: pip install pytesseract pillow && apt install tesseract-ocr]"
        )
    try:
        img = Image.open(path)
        text = pytesseract.image_to_string(img, lang="eng").strip()
        if text:
            logger.info("Image OCR extracted: %d chars from %s", len(text), path.name)
            return text
        return f"[No readable text found in image {path.name} — may be a diagram or low-quality scan]"
    except Exception as exc:
        logger.warning("Image OCR failed for %s: %s", path.name, exc)
        return f"[OCR failed for {path.name}: {exc}]"


def build_extraction_prompt(extracted_text: str, user_message: str) -> str:
    """
    Build a structured extraction prompt for Qwen.
    The model is instructed to return key fields in a fixed format
    so the response can be parsed deterministically.
    """
    return f"""<|im_start|>system
You are AntarAI, a sovereign on-premise AI assistant for MRPL (Mangalore Refinery and Petrochemicals Ltd).
You are analysing an uploaded industrial document. Extract structured information and respond ONLY in the exact format below.
Do not add any preamble, explanation, or extra text outside the format.
<|im_end|>
<|im_start|>user
Extract the following from this inspection/maintenance report. Respond in this EXACT format (fill in <...> placeholders):

Inspection ID: <id or N/A>
Equipment: <equipment name/tag or N/A>
Inspection Date: <date or N/A>
Finding: <one sentence summary of the key finding>
Severity: <Low / Medium / High / Critical>
Recommended Action: <specific recommended action>
Summary: <2-3 sentence overall summary of the document>

Document text:
{extracted_text[:3000]}

User query: {user_message}
<|im_end|>
<|im_start|>assistant
"""


def parse_extraction_response(response: str) -> dict:
    """
    Parse the structured Qwen response into a dict of fields.
    Robust to slight formatting variations.
    """
    fields = {
        "inspection_id": "N/A",
        "equipment": "N/A",
        "inspection_date": "N/A",
        "finding": "N/A",
        "severity": "N/A",
        "recommended_action": "N/A",
        "summary": "",
    }
    key_map = {
        "inspection id": "inspection_id",
        "equipment": "equipment",
        "inspection date": "inspection_date",
        "finding": "finding",
        "severity": "severity",
        "recommended action": "recommended_action",
        "summary": "summary",
    }
    for line in response.splitlines():
        if ":" in line:
            raw_key, _, value = line.partition(":")
            key = raw_key.strip().lower()
            if key in key_map:
                fields[key_map[key]] = value.strip()
    return fields
